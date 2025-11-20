# ABOUTME: Generates test fixture files (PDF, DOCX) for testing conversion.
# ABOUTME: Creates simple documents with known content for validation.

from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def generate_test_docx(output_path: Path):
    """Generate a simple test DOCX file."""
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph.')
    doc.add_paragraph('This is another paragraph with some content.')
    doc.save(output_path)


def generate_test_pdf(output_path: Path):
    """Generate a simple test PDF file."""
    c = canvas.Canvas(str(output_path), pagesize=letter)
    c.drawString(100, 750, "Test PDF Document")
    c.drawString(100, 730, "This is a test paragraph.")
    c.drawString(100, 710, "This is another paragraph with some content.")
    c.save()


if __name__ == '__main__':
    fixtures_dir = Path(__file__).parent / 'fixtures'
    fixtures_dir.mkdir(exist_ok=True)

    generate_test_docx(fixtures_dir / 'test_document.docx')
    generate_test_pdf(fixtures_dir / 'test_document.pdf')

    print(f"Generated test fixtures in {fixtures_dir}")
